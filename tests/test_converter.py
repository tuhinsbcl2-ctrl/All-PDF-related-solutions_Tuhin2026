"""
Tests for core/converter.py
"""

import os
import pytest
import tempfile
from pathlib import Path

# ---------------------------------------------------------------------------
# Helpers – create minimal test fixtures without needing real PDFs
# ---------------------------------------------------------------------------

def _make_test_pdf(tmp_path: Path, text: str = "Hello PDF") -> str:
    """Create a minimal single-page PDF using reportlab and return path."""
    from reportlab.pdfgen import canvas
    path = str(tmp_path / "test.pdf")
    c = canvas.Canvas(path)
    c.drawString(72, 700, text)
    c.save()
    return path


def _make_test_excel(tmp_path: Path) -> str:
    """Create a minimal Excel file and return path."""
    import openpyxl
    path = str(tmp_path / "test.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Name", "Value"])
    ws.append(["Alpha", 1])
    ws.append(["Beta", 2])
    wb.save(path)
    return path


def _make_test_docx(tmp_path: Path) -> str:
    """Create a minimal .docx file and return path."""
    from docx import Document
    path = str(tmp_path / "test.docx")
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.save(path)
    return path


def _make_test_image(tmp_path: Path) -> str:
    """Create a minimal PNG image and return path."""
    from PIL import Image
    path = str(tmp_path / "test.png")
    img = Image.new("RGB", (100, 100), color=(255, 0, 0))
    img.save(path)
    return path


# ---------------------------------------------------------------------------
# pdf_to_excel
# ---------------------------------------------------------------------------

def test_pdf_to_excel_creates_file(tmp_path):
    from src.core.converter import pdf_to_excel
    pdf = _make_test_pdf(tmp_path, "Name Value\nAlpha 1\nBeta 2")
    out = str(tmp_path / "output.xlsx")
    result = pdf_to_excel(pdf, out)
    assert result == out
    assert os.path.isfile(out)


def test_pdf_to_excel_file_not_found(tmp_path):
    from src.core.converter import pdf_to_excel
    with pytest.raises(FileNotFoundError):
        pdf_to_excel("/nonexistent/file.pdf", str(tmp_path / "out.xlsx"))


# ---------------------------------------------------------------------------
# excel_to_pdf
# ---------------------------------------------------------------------------

def test_excel_to_pdf_creates_file(tmp_path):
    from src.core.converter import excel_to_pdf
    xlsx = _make_test_excel(tmp_path)
    out = str(tmp_path / "output.pdf")
    result = excel_to_pdf(xlsx, out)
    assert result == out
    assert os.path.isfile(out)
    assert os.path.getsize(out) > 0


def test_excel_to_pdf_file_not_found(tmp_path):
    from src.core.converter import excel_to_pdf
    with pytest.raises(FileNotFoundError):
        excel_to_pdf("/nonexistent/file.xlsx", str(tmp_path / "out.pdf"))


# ---------------------------------------------------------------------------
# word_to_pdf
# ---------------------------------------------------------------------------

def test_word_to_pdf_creates_file(tmp_path):
    from src.core.converter import word_to_pdf
    docx = _make_test_docx(tmp_path)
    out = str(tmp_path / "output.pdf")
    result = word_to_pdf(docx, out)
    assert result == out
    assert os.path.isfile(out)
    assert os.path.getsize(out) > 0


def test_word_to_pdf_file_not_found(tmp_path):
    from src.core.converter import word_to_pdf
    with pytest.raises(FileNotFoundError):
        word_to_pdf("/nonexistent/file.docx", str(tmp_path / "out.pdf"))


# ---------------------------------------------------------------------------
# pdf_to_word
# ---------------------------------------------------------------------------

def test_pdf_to_word_creates_file(tmp_path):
    from src.core.converter import pdf_to_word
    pdf = _make_test_pdf(tmp_path, "Sample text content")
    out = str(tmp_path / "output.docx")
    result = pdf_to_word(pdf, out)
    assert result == out
    assert os.path.isfile(out)
    assert os.path.getsize(out) > 0


def test_pdf_to_word_file_not_found(tmp_path):
    from src.core.converter import pdf_to_word
    with pytest.raises(FileNotFoundError):
        pdf_to_word("/nonexistent/file.pdf", str(tmp_path / "out.docx"))


# ---------------------------------------------------------------------------
# photos_to_pdf
# ---------------------------------------------------------------------------

def test_photos_to_pdf_creates_file(tmp_path):
    from src.core.converter import photos_to_pdf
    img = _make_test_image(tmp_path)
    out = str(tmp_path / "output.pdf")
    result = photos_to_pdf([img], out)
    assert result == out
    assert os.path.isfile(out)
    assert os.path.getsize(out) > 0


def test_photos_to_pdf_multiple_images(tmp_path):
    from src.core.converter import photos_to_pdf
    from PIL import Image
    imgs = []
    for i in range(3):
        p = str(tmp_path / f"img{i}.png")
        Image.new("RGB", (80, 80), color=(i * 80, 0, 0)).save(p)
        imgs.append(p)
    out = str(tmp_path / "output.pdf")
    result = photos_to_pdf(imgs, out)
    assert os.path.isfile(result)


def test_photos_to_pdf_no_images(tmp_path):
    from src.core.converter import photos_to_pdf
    with pytest.raises(ValueError):
        photos_to_pdf([], str(tmp_path / "out.pdf"))


# ---------------------------------------------------------------------------
# pdf_to_photos
# ---------------------------------------------------------------------------

def test_pdf_to_photos_creates_images(tmp_path):
    from src.core.converter import pdf_to_photos
    pdf = _make_test_pdf(tmp_path)
    out_dir = str(tmp_path / "images")
    result = pdf_to_photos(pdf, out_dir, dpi=72)
    assert len(result) >= 1
    for p in result:
        assert os.path.isfile(p)


def test_pdf_to_photos_file_not_found(tmp_path):
    from src.core.converter import pdf_to_photos
    with pytest.raises(FileNotFoundError):
        pdf_to_photos("/nonexistent/file.pdf", str(tmp_path / "images"))
