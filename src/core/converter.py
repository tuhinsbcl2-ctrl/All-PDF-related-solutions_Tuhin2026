"""
Conversion operations:
  - PDF ↔ Excel
  - PDF ↔ Word
  - Photo ↔ PDF
"""

import os
from pathlib import Path

import fitz  # PyMuPDF
import openpyxl
from openpyxl.styles import Font, Alignment
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
)
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image
import docx
from docx import Document


# ---------------------------------------------------------------------------
# PDF → Excel
# ---------------------------------------------------------------------------

def pdf_to_excel(pdf_path: str, output_path: str) -> str:
    """
    Extract text content from every page of a PDF and write it to an Excel file.
    Each page becomes a separate worksheet.

    Returns the output path.
    """
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    doc = fitz.open(pdf_path)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    for page_index in range(len(doc)):
        page = doc[page_index]
        sheet_name = f"Page {page_index + 1}"
        ws = wb.create_sheet(title=sheet_name)

        # Extract text blocks
        blocks = page.get_text("blocks")  # list of (x0,y0,x1,y1,text,block_no,block_type)
        row = 1
        for block in blocks:
            if block[6] == 0:  # text block
                text = block[4].strip()
                if text:
                    lines = text.split("\n")
                    for line in lines:
                        ws.cell(row=row, column=1, value=line)
                        ws.cell(row=row, column=1).font = Font(name="Calibri", size=10)
                        row += 1

        # Auto-size column
        ws.column_dimensions["A"].width = 80

    doc.close()
    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Excel → PDF
# ---------------------------------------------------------------------------

def excel_to_pdf(excel_path: str, output_path: str) -> str:
    """
    Read an Excel file and generate a PDF with tables.

    Returns the output path.
    """
    if not os.path.isfile(excel_path):
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    styles = getSampleStyleSheet()
    story = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        story.append(Paragraph(f"Sheet: {sheet_name}", styles["Heading2"]))
        story.append(Spacer(1, 6))

        data = []
        for row in ws.iter_rows(values_only=True):
            data.append([str(cell) if cell is not None else "" for cell in row])

        if data:
            col_count = max(len(r) for r in data)
            col_width = (A4[0] - 72) / max(col_count, 1)
            table = Table(data, colWidths=[col_width] * col_count)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2196F3")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("WORDWRAP", (0, 0), (-1, -1), True),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 12))

    doc.build(story)
    return output_path


# ---------------------------------------------------------------------------
# Word → PDF
# ---------------------------------------------------------------------------

def word_to_pdf(docx_path: str, output_path: str) -> str:
    """
    Read a .docx file and render its text content to a PDF.

    Returns the output path.
    """
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"Word document not found: {docx_path}")

    document = Document(docx_path)
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72,
    )
    styles = getSampleStyleSheet()
    story = []

    for para in document.paragraphs:
        text = para.text
        if not text.strip():
            story.append(Spacer(1, 6))
            continue

        style_name = "Normal"
        if para.style.name.startswith("Heading 1"):
            style_name = "Heading1"
        elif para.style.name.startswith("Heading 2"):
            style_name = "Heading2"
        elif para.style.name.startswith("Heading 3"):
            style_name = "Heading3"

        story.append(Paragraph(text, styles[style_name]))
        story.append(Spacer(1, 4))

    for table in document.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        if data:
            col_count = max(len(r) for r in data)
            col_width = (A4[0] - 144) / max(col_count, 1)
            rl_table = Table(data, colWidths=[col_width] * col_count)
            rl_table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]
                )
            )
            story.append(rl_table)
            story.append(Spacer(1, 8))

    if not story:
        story.append(Paragraph("(empty document)", styles["Normal"]))

    doc.build(story)
    return output_path


# ---------------------------------------------------------------------------
# PDF → Word
# ---------------------------------------------------------------------------

def pdf_to_word(pdf_path: str, output_path: str) -> str:
    """
    Extract text from a PDF and write it to a .docx file.

    Returns the output path.
    """
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    pdf_doc = fitz.open(pdf_path)
    word_doc = Document()
    word_doc.add_heading("Converted from PDF", level=1)

    for page_index in range(len(pdf_doc)):
        page = pdf_doc[page_index]
        text = page.get_text()
        word_doc.add_heading(f"Page {page_index + 1}", level=2)
        for line in text.split("\n"):
            if line.strip():
                word_doc.add_paragraph(line)

    pdf_doc.close()
    word_doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Photos → PDF
# ---------------------------------------------------------------------------

def photos_to_pdf(image_paths: list, output_path: str) -> str:
    """
    Convert one or more image files (PNG, JPG, BMP, TIFF) to a single PDF.
    Each image becomes one page.

    Returns the output path.
    """
    if not image_paths:
        raise ValueError("No image paths provided.")

    images = []
    for path in image_paths:
        if not os.path.isfile(path):
            raise FileNotFoundError(f"Image not found: {path}")
        img = Image.open(path).convert("RGB")
        images.append(img)

    first = images[0]
    rest = images[1:] if len(images) > 1 else []
    first.save(output_path, "PDF", save_all=True, append_images=rest, resolution=150)
    return output_path


# ---------------------------------------------------------------------------
# PDF → Photos
# ---------------------------------------------------------------------------

def pdf_to_photos(pdf_path: str, output_dir: str, dpi: int = 150, fmt: str = "PNG") -> list:
    """
    Render each page of a PDF as an image file.

    Returns list of output image file paths.
    """
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    os.makedirs(output_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    output_paths = []
    ext = fmt.lower()
    if ext == "jpeg":
        ext = "jpg"

    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat, alpha=False)
        stem = Path(pdf_path).stem
        out_file = os.path.join(output_dir, f"{stem}_page_{i + 1}.{ext}")
        pix.save(out_file)
        output_paths.append(out_file)

    doc.close()
    return output_paths
